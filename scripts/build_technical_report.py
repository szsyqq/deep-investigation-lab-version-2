from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "product-technical-report.md"
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT = OUTPUT_DIR / "Deep_Investigation_Lab_V2_产品技术报告.docx"

INK = "17201C"
RED = "B4362F"
GOLD = "A67C2D"
MUTED = "687069"
LIGHT = "EEF0EB"
PALE_RED = "F7EAE7"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None, east_asia="Noto Sans CJK SC"):
    run.font.name = "Noto Sans CJK SC"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_bottom_border(paragraph, color=RED, size="18", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 30, INK, 0, 8),
        ("Subtitle", 14, MUTED, 0, 14),
        ("Heading 1", 18, INK, 18, 8),
        ("Heading 2", 14, RED, 14, 6),
        ("Heading 3", 11.5, INK, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans CJK SC")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_header_footer(section):
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("DEEP INVESTIGATION LAB V2  /  PRODUCT TECHNICAL REPORT")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    add_bottom_border(header, color="C9CEC8", size="6", space="4")
    footer = section.footer.paragraphs[0]
    add_page_field(footer)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    add_header_footer(section)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("PRODUCT TECHNICAL REPORT")
    set_run_font(r, size=10, bold=True, color=RED)

    title = doc.add_paragraph(style="Title")
    title.paragraph_format.space_before = Pt(22)
    title.add_run("Deep Investigation Lab V2")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("从开放式问题到可核查调查报告的\n系统原理、生产流程与工程治理")
    add_bottom_border(subtitle, color=RED, size="20", space="12")

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(30)
    summary.paragraph_format.space_after = Pt(12)
    summary.paragraph_format.line_spacing = 1.35
    r = summary.add_run(
        "面向产品经理、内容产品负责人、研究负责人和协作开发者的系统说明书。"
        "本报告解释每一项中间产物为什么存在、如何形成、如何通过质量关卡，以及在长期维护中需要保护的边界。"
    )
    set_run_font(r, size=12, color=INK)

    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("文档版本", "2.0"),
        ("系统版本", "Deep Investigation Lab V2"),
        ("更新日期", "2026-07-25"),
        ("公开站点", "szsyqq.github.io/deep-investigation-lab-version-2"),
    ]
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        set_cell_shading(table.cell(i, 0), LIGHT)
        for run in table.cell(i, 0).paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=INK)
        for run in table.cell(i, 1).paragraphs[0].runs:
            set_run_font(run, size=9.5, color=MUTED)
    set_table_geometry(table, [1900, 7460])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("调查研究不是一次生成，而是一系列可观察、可核查、可回退的产品状态。")
    set_run_font(r, size=11, italic=True, color=GOLD)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_document_map(doc):
    p = doc.add_paragraph()
    r = p.add_run("阅读导航")
    set_run_font(r, size=10, bold=True, color=RED)
    title = doc.add_paragraph("本报告如何组织", style="Heading 1")
    title.paragraph_format.space_before = Pt(4)
    items = [
        ("01—03", "定位、问题重构与总体架构"),
        ("04—05", "研究任务识别与六道质量关卡"),
        ("06—09", "修改机制、发布工程、日常操作与治理"),
        ("10—12", "仓库架构、知识更新与模型优化"),
        ("13—15", "风险处理、后续路线与结论"),
        ("附录", "研究包职责与常用命令"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "章节"
    table.cell(0, 1).text = "内容"
    set_cell_shading(table.cell(0, 0), INK)
    set_cell_shading(table.cell(0, 1), INK)
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=WHITE)
    for chapter, content in items:
        cells = table.add_row().cells
        cells[0].text = chapter
        cells[1].text = content
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=RED)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1900, 7460])

    doc.add_paragraph("系统主链路", style="Heading 2")
    flow = doc.add_table(rows=1, cols=5)
    labels = ["研究请求", "研究包", "编辑判断", "发布合同", "共享网站"]
    fills = [RED, INK, GOLD, INK, RED]
    for idx, label in enumerate(labels):
        cell = flow.cell(0, idx)
        cell.text = label
        set_cell_shading(cell, fills[idx])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.5, bold=True, color=WHITE)
    set_table_geometry(flow, [1872] * 5)

    callout = doc.add_table(rows=1, cols=1)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_RED)
    p = cell.paragraphs[0]
    r = p.add_run("核心判断  ")
    set_run_font(r, size=10, bold=True, color=RED)
    r = p.add_run("V2 的创新不是增加更多提示词，而是让每一次研究与编辑决策形成可检查的中间状态。")
    set_run_font(r, size=10.5, color=INK)
    set_table_geometry(callout, [PAGE_WIDTH_DXA])
    doc.add_page_break()


def add_inline_markup(paragraph, text):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=9.5, color=RED)
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, color=INK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_blockquote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_RED)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, italic=True, color=INK)
    set_table_geometry(table, [PAGE_WIDTH_DXA])


def parse_source(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    skip_front = True
    paragraph_buffer = []
    break_before = {
        "3. 系统总体架构",
        "5. 六道研究关卡的详细流程",
        "8. 面向产品经理的日常操作流程",
        "10. GitHub 仓库架构与文件维护",
        "14. 后续开发路线",
        "附录 A：研究包文件职责",
    }

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_inline_markup(p, " ".join(paragraph_buffer))
            paragraph_buffer = []

    for raw in lines:
        line = raw.strip()
        if skip_front:
            if line == "## 执行摘要":
                skip_front = False
            else:
                continue
        if not line:
            flush()
            continue
        if line == "---":
            flush()
            continue
        if line.startswith("## "):
            flush()
            title = line[3:]
            is_subsection = bool(re.match(r"^\d+\.\d+(?:\s|$)", title))
            if title in break_before and not is_subsection:
                doc.add_page_break()
            doc.add_paragraph(title, style="Heading 2" if is_subsection else "Heading 1")
        elif line.startswith("### "):
            flush()
            title = line[4:]
            is_numbered = bool(re.match(r"^\d+\.\d+(?:\s|$)", title))
            doc.add_paragraph(title, style="Heading 2" if is_numbered else "Heading 3")
        elif line.startswith("#### "):
            flush()
            doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("> "):
            flush()
            add_blockquote(doc, line[2:])
        elif re.match(r"^\\d+\\.\\s+", line):
            flush()
            text = re.sub(r"^\\d+\\.\\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_markup(p, text)
        elif line.startswith("- "):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, line[2:])
        else:
            paragraph_buffer.append(line)
    flush()


def add_end_matter(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("END OF REPORT")
    set_run_font(r, size=9, bold=True, color=RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Deep Investigation Lab V2")
    set_run_font(r, size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("公开站点  ·  GitHub 仓库  ·  操作手册")
    set_run_font(r, size=10, color=MUTED)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_cover(doc)
    add_document_map(doc)
    parse_source(doc)
    add_end_matter(doc)
    doc.core_properties.title = "Deep Investigation Lab V2 产品技术报告"
    doc.core_properties.subject = "系统原理、生产流程与工程治理"
    doc.core_properties.author = "Deep Investigation Lab"
    doc.core_properties.keywords = "调查研究, 产品架构, 编辑工作流, 视觉规划, GitHub Pages"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
